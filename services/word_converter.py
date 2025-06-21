#!/usr/bin/env python3
"""
Word Document to Markdown Converter Service

This service handles conversion of Word documents (.docx, .doc) to Markdown format.
"""

import sys
from pathlib import Path
from typing import List
import re
import zipfile
import io

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is not installed. Please install it using:")
    print("pip install python-docx")
    sys.exit(1)

from .base_converter import BaseDocumentConverter


class WordDocumentConverter(BaseDocumentConverter):
    """Converts Word documents to Markdown format with full content preservation."""

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported Word document extensions."""
        return ['.docx', '.doc']
    
    def can_convert(self, file_path: Path) -> bool:
        """Check if this converter can handle Word documents."""
        return file_path.suffix.lower() in self.get_supported_extensions()

    def _clean_formatting_markers(self, text: str) -> str:
        """Clean up consecutive formatting markers in text."""
        # Fix consecutive bold markers: **text****more** -> **text more**
        text = re.sub(r'\*\*([^*]*)\*\*\*\*([^*]*)\*\*', r'**\1\2**', text)

        # Fix consecutive italic markers: *text**more* -> *text more*
        text = re.sub(r'\*([^*]*)\*\*([^*]*)\*', r'*\1\2*', text)

        # Fix mixed markers: **bold***italic* -> ***bold italic***
        text = re.sub(r'\*\*([^*]*)\*\*\*([^*]*)\*', r'***\1\2***', text)

        # Remove empty formatting markers
        text = re.sub(r'\*\*\*\*', '', text)  # Empty bold
        text = re.sub(r'(?<!\*)\*\*(?!\*)', '', text)  # Empty bold (not part of italic)
        text = re.sub(r'(?<!\*)\*(?!\*)', '', text)  # Empty italic

        return text
    
    def _convert_paragraph_to_markdown(self, paragraph: Paragraph) -> str:
        """Convert a Word paragraph to Markdown format."""
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # Handle different paragraph styles
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        # Check if this looks like a heading based on style or content
        is_heading = False
        heading_level = 1
        
        # First check style-based headings
        if 'heading 1' in style_name:
            is_heading = True
            heading_level = 1
        elif 'heading 2' in style_name:
            is_heading = True
            heading_level = 2
        elif 'heading 3' in style_name:
            is_heading = True
            heading_level = 3
        elif 'heading 4' in style_name:
            is_heading = True
            heading_level = 4
        elif 'heading 5' in style_name:
            is_heading = True
            heading_level = 5
        elif 'heading 6' in style_name:
            is_heading = True
            heading_level = 6
        else:
            # Only detect headings if they have clear indicators to preserve original content
            # Check for explicit section numbering patterns first
            section_patterns = [
                r'^(\d+(?:\.\d+)*)\s+(.+)$',  # "4.10 Glossary", "1.2.3 Title" (with space)
            ]

            section_match = None
            for pattern in section_patterns:
                match = re.match(pattern, text)
                if match:
                    section_match = match
                    break

            if section_match:
                # Determine heading level based on number of dots
                section_number = section_match.group(1)
                dot_count = section_number.count('.')
                heading_level = min(dot_count + 1, 6)  # Cap at level 6
                is_heading = True
            else:
                # Be more conservative - only detect as heading if it's clearly formatted as one
                # Check if paragraph is short, all bold, and looks like a title
                if len(text) < 100 and len(text.split()) <= 10:  # Short and concise
                    all_bold = True
                    has_text = False

                    for run in paragraph.runs:
                        if run.text.strip():
                            has_text = True
                            if not run.bold:
                                all_bold = False
                                break

                    # Only treat as heading if it's all bold AND matches common heading patterns
                    if all_bold and has_text:
                        # Check for common document section patterns
                        heading_patterns = [
                            r'^[A-Z][A-Z\s&/]+$',  # All caps like "INTRODUCTION", "BUSINESS REQUIREMENTS"
                            r'^\d+\.\s*[A-Z]',     # Numbered sections like "1. Introduction"
                            r'^[A-Z][a-z]+(\s+[A-Z][a-z]*)*:?$',  # Title case like "Business Requirements"
                        ]

                        is_likely_heading = any(re.match(pattern, text) for pattern in heading_patterns)

                        if is_likely_heading:
                            is_heading = True
                            heading_level = 2  # Default to level 2 for detected headings

        # Convert to markdown heading if identified as heading
        if is_heading:
            heading_prefix = "#" * heading_level
            section_number = self._update_section_counter(heading_level)
            return f"{heading_prefix} {section_number}{text}\n\n"

        # Handle formatting within runs for regular paragraphs
        markdown_text = ""
        for run in paragraph.runs:
            run_text = run.text

            # Apply bold formatting
            if run.bold and run_text.strip():
                run_text = f"**{run_text}**"

            # Apply italic formatting
            if run.italic and run_text.strip():
                run_text = f"*{run_text}*"

            markdown_text += run_text

        # Clean up multiple consecutive formatting markers
        markdown_text = self._clean_formatting_markers(markdown_text)

        return f"{markdown_text}\n\n"
    
    def _extract_cell_content(self, cell) -> str:
        """Extract content from a table cell, handling multiple paragraphs properly."""
        if not cell.paragraphs:
            return ""

        # Join all paragraphs in the cell with a space to keep content in one line
        # This prevents table structure from breaking due to newlines
        cell_parts = []
        for paragraph in cell.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                cell_parts.append(para_text)

        # Join with space and clean up any extra whitespace
        cell_content = " ".join(cell_parts).strip()

        # Replace any remaining newlines with spaces to ensure table integrity
        cell_content = cell_content.replace('\n', ' ').replace('\r', ' ')

        # Clean up multiple consecutive spaces
        cell_content = ' '.join(cell_content.split())

        return cell_content

    def _convert_table_to_markdown(self, table: Table) -> str:
        """Convert a Word table to Markdown table format."""
        if not table.rows:
            return ""

        markdown_table = ""

        # Process header row
        header_row = table.rows[0]
        header_cells = [self._extract_cell_content(cell) for cell in header_row.cells]
        markdown_table += "| " + " | ".join(header_cells) + " |\n"

        # Add separator row
        separator = "| " + " | ".join(["---"] * len(header_cells)) + " |\n"
        markdown_table += separator

        # Process data rows
        for row in table.rows[1:]:
            row_cells = [self._extract_cell_content(cell) for cell in row.cells]
            markdown_table += "| " + " | ".join(row_cells) + " |\n"

        return markdown_table + "\n"
    
    def _analyze_document_structure(self, doc: Document) -> None:
        """Analyze document structure to understand heading patterns."""
        self.logger.info("Analyzing Word document structure...")
        
        section_pattern = r'^(\d+(?:\.\d+)*)\s*(.*)$'
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text and len(text) < 200:  # Focus on potential headings
                style_name = paragraph.style.name if paragraph.style else "None"
                
                # Check for section numbering
                match = re.match(section_pattern, text)
                if match:
                    section_num = match.group(1)
                    section_title = match.group(2).strip()
                    self.logger.info(f"Found section: '{section_num}' - '{section_title}' (Style: {style_name})")
                
                # Check for bold formatting
                all_bold = True
                has_text = False
                for run in paragraph.runs:
                    if run.text.strip():
                        has_text = True
                        if not run.bold:
                            all_bold = False
                            break
                
                if all_bold and has_text and 'heading' not in style_name.lower():
                    self.logger.info(f"Potential bold heading: '{text[:50]}...' (Style: {style_name})")

    def _extract_images_from_word_document(self, doc_path: Path) -> List[Path]:
        """
        Extract embedded images from Word document and save them as temporary files.

        Args:
            doc_path: Path to the Word document

        Returns:
            List of paths to extracted image files
        """
        extracted_images = []

        if not self.extract_images:
            return extracted_images

        try:
            # Word documents are ZIP files, extract images from the media folder
            with zipfile.ZipFile(doc_path, 'r') as docx_zip:
                # List all files in the ZIP
                file_list = docx_zip.namelist()

                # Find image files in word/media/ directory
                image_files = [f for f in file_list if f.startswith('word/media/') and
                             any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.svg'])]

                if image_files:
                    # Create temp directory for images
                    temp_dir = self._create_temp_image_dir()

                    for img_file in image_files:
                        try:
                            # Extract image data
                            img_data = docx_zip.read(img_file)

                            # Create filename from the original path
                            img_filename = Path(img_file).name
                            temp_img_path = temp_dir / img_filename

                            # Write image to temp file
                            with open(temp_img_path, 'wb') as f:
                                f.write(img_data)

                            extracted_images.append(temp_img_path)
                            self.logger.info(f"Extracted image: {img_filename}")

                        except Exception as e:
                            self.logger.warning(f"Failed to extract image {img_file}: {str(e)}")

                self.logger.info(f"Extracted {len(extracted_images)} images from {doc_path.name}")

        except Exception as e:
            self.logger.error(f"Failed to extract images from {doc_path}: {str(e)}")

        return extracted_images

    def _convert_document_to_markdown(self, doc_path: Path) -> str:
        """Convert a Word document to Markdown format."""
        try:
            doc = Document(doc_path)

            # Reset section counters for new document
            self._reset_section_counters()

            # Analyze document structure for debugging
            self._analyze_document_structure(doc)

            # Extract embedded images first
            extracted_images = self._extract_images_from_word_document(doc_path)

            markdown_content = ""

            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find the corresponding paragraph object
                    for paragraph in doc.paragraphs:
                        if paragraph._element == element:
                            markdown_content += self._convert_paragraph_to_markdown(paragraph)
                            break

                elif element.tag.endswith('tbl'):  # Table
                    # Find the corresponding table object
                    for table in doc.tables:
                        if table._element == element:
                            markdown_content += self._convert_table_to_markdown(table)
                            break

            # Add extracted image content at the end of the document
            if extracted_images:
                markdown_content += "\n\n# Embedded Images\n\n"
                markdown_content += "The following content was extracted from embedded images in the document:\n\n"

                for img_path in extracted_images:
                    image_markdown = self._convert_image_to_markdown(img_path)
                    markdown_content += image_markdown

            return markdown_content

        except Exception as e:
            self.logger.error(f"Error converting Word document {doc_path}: {str(e)}")
            return f"# Error Converting Document\n\nFailed to convert {doc_path.name}: {str(e)}\n"
