#!/usr/bin/env python3
"""
Test script for the Document to Markdown Converter

This script creates a sample Word document for testing purposes.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Error: python-docx is not installed. Please install it using:")
    print("pip install python-docx")
    sys.exit(1)

from document_converter import DocumentToMarkdownConverter


def create_sample_document():
    """Create a sample Word document for testing."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Document for Testing', 0)
    
    # Add some paragraphs
    doc.add_heading('Introduction', level=1)
    p1 = doc.add_paragraph('This is a ')
    p1.add_run('sample document').bold = True
    p1.add_run(' created for testing the ')
    p1.add_run('Document to Markdown Converter').italic = True
    p1.add_run('.')
    
    doc.add_paragraph('This paragraph contains regular text without any special formatting.')
    
    # Add a second-level heading
    doc.add_heading('Features Demonstrated', level=2)
    
    # Add a list (as paragraphs since docx doesn't have direct list support in python-docx)
    doc.add_paragraph('• Bold text formatting')
    doc.add_paragraph('• Italic text formatting')
    doc.add_paragraph('• Multiple heading levels')
    doc.add_paragraph('• Tables (see below)')
    
    # Add a table
    doc.add_heading('Sample Table', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Age'
    hdr_cells[2].text = 'City'
    
    # Add data rows
    row_data = [
        ('Alice', '25', 'New York'),
        ('Bob', '30', 'San Francisco'),
        ('Charlie', '35', 'Chicago')
    ]
    
    for name, age, city in row_data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = age
        row_cells[2].text = city
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion = doc.add_paragraph('This document demonstrates various formatting features that can be converted to Markdown. ')
    conclusion.add_run('The converter should handle').bold = True
    conclusion.add_run(' all these elements properly.')
    
    # Save the document
    input_folder = Path('input')
    input_folder.mkdir(exist_ok=True)
    doc_path = input_folder / 'sample_document.docx'
    doc.save(doc_path)
    
    print(f"Sample document created: {doc_path}")
    return doc_path


def test_conversion():
    """Test the conversion process."""
    print("Document to Markdown Converter - Test")
    print("=" * 40)
    
    # Create sample document
    sample_doc = create_sample_document()
    
    # Initialize converter
    converter = DocumentToMarkdownConverter()
    
    # Convert the sample document
    print("\nTesting conversion...")
    success = converter.convert_file(sample_doc)
    
    if success:
        print("✅ Test conversion successful!")
        
        # Show the converted content
        output_file = Path('output') / 'sample_document.md'
        if output_file.exists():
            print(f"\nConverted content preview:")
            print("-" * 40)
            with open(output_file, 'r', encoding='utf-8') as f:
                content = f.read()
                # Show first 500 characters
                preview = content[:500]
                print(preview)
                if len(content) > 500:
                    print("...")
                    print(f"\n[Full content saved to: {output_file}]")
    else:
        print("❌ Test conversion failed!")


if __name__ == "__main__":
    test_conversion()
